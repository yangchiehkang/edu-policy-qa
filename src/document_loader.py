import json
import re
from pathlib import Path

import pandas as pd
import fitz
from docx import Document
from tqdm import tqdm


PROJECT_ROOT = Path(r"E:\Working\edu_policy_qa")

RAW_DOCS_DIR = PROJECT_ROOT / "data" / "raw_docs"
METADATA_FILE = PROJECT_ROOT / "data" / "metadata" / "raw_docs_metadata_clean.csv"
OUTPUT_FILE = PROJECT_ROOT / "data" / "processed_docs" / "documents.jsonl"
ERROR_FILE = PROJECT_ROOT / "data" / "processed_docs" / "document_parse_errors.csv"


def clean_text(text: str) -> str:
    if text is None:
        return ""

    text = text.replace("\x00", " ")
    text = text.replace("\ufeff", " ")
    text = text.replace("\r", "\n")

    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)

    return text.strip()


def read_txt(file_path: Path):
    encodings = ["utf-8", "utf-8-sig", "latin-1", "cp1252"]

    for enc in encodings:
        try:
            text = file_path.read_text(encoding=enc)
            text = clean_text(text)

            if text:
                return [
                    {
                        "page": 1,
                        "text": text,
                    }
                ]
        except Exception:
            continue

    return []


def read_pdf(file_path: Path):
    pages = []

    doc = fitz.open(file_path)

    for page_idx in range(len(doc)):
        page = doc.load_page(page_idx)
        text = page.get_text("text")
        text = clean_text(text)

        if text:
            pages.append(
                {
                    "page": page_idx + 1,
                    "text": text,
                }
            )

    doc.close()
    return pages


def read_docx(file_path: Path):
    doc = Document(file_path)

    paragraphs = []
    for para in doc.paragraphs:
        text = clean_text(para.text)
        if text:
            paragraphs.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = clean_text(cell.text)
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                paragraphs.append(" | ".join(row_text))

    full_text = clean_text("\n".join(paragraphs))

    if not full_text:
        return []

    return [
        {
            "page": 1,
            "text": full_text,
        }
    ]


def read_md_or_html(file_path: Path):
    return read_txt(file_path)


def parse_file(file_path: Path):
    suffix = file_path.suffix.lower()

    if suffix == ".txt":
        return read_txt(file_path)

    if suffix == ".pdf":
        return read_pdf(file_path)

    if suffix == ".docx":
        return read_docx(file_path)

    if suffix in [".md", ".html", ".htm"]:
        return read_md_or_html(file_path)

    return []


def main():
    OUTPUT_FILE.parent.mkdir(parents=True, exist_ok=True)

    metadata = pd.read_csv(METADATA_FILE)

    records = []
    errors = []

    for _, row in tqdm(metadata.iterrows(), total=len(metadata), desc="Parsing documents"):
        doc_id = str(row["doc_id"])
        file_name = str(row["file_name"])
        file_path = RAW_DOCS_DIR / file_name

        if not file_path.exists():
            errors.append(
                {
                    "doc_id": doc_id,
                    "file_name": file_name,
                    "error": "file not found",
                }
            )
            continue

        try:
            pages = parse_file(file_path)

            if not pages:
                errors.append(
                    {
                        "doc_id": doc_id,
                        "file_name": file_name,
                        "error": "empty parsed text",
                    }
                )
                continue

            for item in pages:
                text = clean_text(item["text"])

                if not text:
                    continue

                records.append(
                    {
                        "doc_id": doc_id,
                        "source": file_name,
                        "doc_type": str(row.get("doc_type", "")),
                        "publication_date": str(row.get("publication_date", "")),
                        "data_source": str(row.get("source", "")),
                        "source_url": str(row.get("source_url", "")),
                        "is_official_course_material": str(row.get("is_official_course_material", "")),
                        "page": int(item["page"]),
                        "text": text,
                    }
                )

        except Exception as e:
            errors.append(
                {
                    "doc_id": doc_id,
                    "file_name": file_name,
                    "error": repr(e),
                }
            )

    with open(OUTPUT_FILE, "w", encoding="utf-8") as f:
        for record in records:
            f.write(json.dumps(record, ensure_ascii=False) + "\n")

    pd.DataFrame(errors).to_csv(ERROR_FILE, index=False, encoding="utf-8-sig")

    print(f"Documents in metadata: {len(metadata)}")
    print(f"Parsed records: {len(records)}")
    print(f"Parse errors: {len(errors)}")
    print(f"Saved documents to: {OUTPUT_FILE}")
    print(f"Saved errors to: {ERROR_FILE}")


if __name__ == "__main__":
    main()
